from typing import Dict, Any, List
from io import BytesIO

import pdfplumber
from docx import Document


class SOWParser:
    """
    Advanced SOW parser that extracts:
    - Text content
    - Tables
    - Images (including architecture diagrams)
    - Structured metadata
    """

    # ==========================================================
    # MAIN ENTRY
    # ==========================================================

    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> Dict[str, Any]:
        if filename.lower().endswith(".pdf"):
            return SOWParser._parse_pdf(file_bytes)
        elif filename.lower().endswith(".docx"):
            return SOWParser._parse_docx(file_bytes)
        else:
            raise ValueError("Unsupported file type. Only PDF and DOCX allowed.")

    # ==========================================================
    # PDF PARSER
    # ==========================================================

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> Dict[str, Any]:
        text_content = []
        tables = []
        images = []

        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:

                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)

                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)

                # Extract images (architecture diagrams included)
                for image in page.images:
                    images.append({
                        "x0": image.get("x0"),
                        "top": image.get("top"),
                        "width": image.get("width"),
                        "height": image.get("height")
                    })

        return {
            "text": "\n".join(text_content),
            "tables": tables,
            "images": images,
            "metadata": SOWParser._extract_structured_metadata("\n".join(text_content))
        }

    # ==========================================================
    # DOCX PARSER
    # ==========================================================

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> Dict[str, Any]:
        document = Document(BytesIO(file_bytes))

        text_content = []
        tables = []
        images = []

        # Extract paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())

        # Extract tables
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        # Extract images (including diagrams)
        for rel in document.part._rels:
            rel_obj = document.part._rels[rel]
            if "image" in rel_obj.target_ref:
                images.append({
                    "image_name": rel_obj.target_ref
                })

        full_text = "\n".join(text_content)

        return {
            "text": full_text,
            "tables": tables,
            "images": images,
            "metadata": SOWParser._extract_structured_metadata(full_text)
        }

    # ==========================================================
    # STRUCTURED METADATA EXTRACTION
    # ==========================================================

    @staticmethod
    def _extract_structured_metadata(text: str) -> Dict[str, Any]:
        """
        Basic keyword-driven extractor.
        Can be replaced with AI-based extraction later.
        """

        return {
            "project_name": SOWParser._extract_field(text, "Project Name"),
            "client": SOWParser._extract_field(text, "Client"),
            "domain": SOWParser._extract_field(text, "Domain"),
            "start_date": SOWParser._extract_field(text, "Start Date"),
            "end_date": SOWParser._extract_field(text, "End Date"),
            "objectives": SOWParser._extract_section(text, "Objective"),
            "scope": SOWParser._extract_section(text, "Scope"),
        }

    # ==========================================================
    # FIELD EXTRACTOR
    # ==========================================================

    @staticmethod
    def _extract_field(text: str, keyword: str) -> str:
        for line in text.split("\n"):
            if keyword.lower() in line.lower():
                parts = line.split(":")
                if len(parts) > 1:
                    return parts[1].strip()
        return ""

    # ==========================================================
    # SECTION EXTRACTOR
    # ==========================================================

    @staticmethod
    def _extract_section(text: str, keyword: str) -> str:
        lines = text.split("\n")
        capture = False
        section_lines = []

        for line in lines:
            if keyword.lower() in line.lower():
                capture = True
                continue

            if capture:
                if line.strip() == "":
                    break
                section_lines.append(line.strip())

        return "\n".join(section_lines)